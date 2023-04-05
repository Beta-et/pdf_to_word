import subprocess
import sys
import docx
from pathlib import Path
import tkinter as tk
from cropper import App


def extract_text_from_image(image_file):
    image_file = 'img/' + image_file
    print("path:", image_file)
    text = subprocess.run(["tesseract", image_file, "stdout", "-l", "eng"], capture_output=True, text=True).stdout
    return text


def header_footer(doc, file_name):
    print('header_footer')
    head = ""
    foot = ""
    path_list = Path('cropped').glob(f'*{file_name}*')
    for path in path_list:
        if str(path)[:-4].endswith('top'):
            head = str(path)
        elif str(path)[:-4].endswith('btm'):
            foot = str(path)

    print("head:", head)
    print("foot:", foot)

    if head != "":
        header = doc.sections[0].header
        header.paragraphs[0].add_run().add_picture(head)

    if foot != "":
        footer = doc.sections[0].footer
        footer.paragraphs[0].add_run().add_picture(foot)


def write_text_to_word(text, file_name):
    try:
        doc = docx.Document(f"word/{file_name}.docx")
        doc.add_paragraph(text)
        doc.save(f"word/{file_name}.docx")
        print('saved')
    except Exception:
        docx.Document().save(f"word/{file_name}.docx")
        write_text_to_word(text, file_name)


def append_text_to_word(text, file_name):
    try:
        doc = docx.Document(f"word/{file_name}.docx")
        doc.add_paragraph(text)
        doc.save(f"word/{file_name}.docx")
        print('saved')
    except Exception:
        docx.Document().save(f"word/{file_name}.docx")
        append_text_to_word(text, file_name)


def check_multiple_page(file_name):
    path_list = Path('img').glob(f'*{file_name}*')
    files = []
    for path in path_list:
        path_in_str = str(path)
        files.append(path_in_str[4:])

    if len(files) == 1:
        return 1, files
    else:
        return len(files), files


def cropper(filename):
    root = tk.Tk()
    app = App(root, filename)
    root.mainloop()


def main(file_name):
    name = str(file_name)
    file_count, files = check_multiple_page(file_name)
    files.sort()

    cropper(files[0])

    print("count:", file_count, "files:", files)

    if file_count == 1:
        if files[0].endswith(".jpg") or files[0].endswith(".png"):
            text = extract_text_from_image(files[0])
        else:
            print("Unsupported file format. Please provide a PDF or image file.")
            sys.exit(1)

        # print(text)
        write_text_to_word(text, name)

    else:
        for file in files:
            if file.endswith(".jpg") or file.endswith(".png"):
                text = extract_text_from_image(file)
            else:
                print("Unsupported file format. Please provide a PDF or image file.")
                sys.exit(1)

            # print(text)
            append_text_to_word(text, name)

    doc = docx.Document(f'word/{name}.docx')
    header_footer(doc, name)
    doc.save(f"word/{name}.docx")

